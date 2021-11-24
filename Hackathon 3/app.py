from flask import Flask, render_template, request, redirect, flash, g, make_response
from docx import Document
from pymysql import connect, cursors
import other

app = Flask(__name__)
app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'

app.register_blueprint(other.bp)

db_config = {
    "host": "localhost",
    "user": "root",
    "password": "Tbd!@240992",
    "database": "blog",
    "cursorclass": cursors.DictCursor
}


def connect_db():
    conn = connect(**db_config)
    cur = conn.cursor()
    g.conn = conn
    g.cur = cur


def close_db():
    conn = g.get("conn")
    cur = g.get("cur")

    if conn and cur:
        cur.close()
        conn.close()


@app.before_request
def before_request():
    connect_db()
    user_id = request.cookies.get('user_id')

    if user_id:
        g.cur.execute("SELECT * FROM users WHERE id = %s", user_id)
        g.user = g.cur.fetchone()


@app.teardown_request
def teardown_request(Exeption):
    close_db()


@app.get("/signup")
def render_signup_page():
    if g.get("user"):
        return redirect("/")
    return render_template("signup.html")


@app.post("/signup")
def signup():
    username = request.form.get("username")
    password = request.form.get("password")
    repassword = request.form.get("repassword")

    if not username or not password or not repassword:
        flash("Username, password and repassword cannot be empty")
    elif len(username.strip()) < 4:
        flash("Username must have at least 4 characters")
    elif len(password) < 6:
        flash("Password must have at least 6 characters")
    elif password != repassword:
        flash("Re-Password is not correct")
    else:
        sql = '''
            SELECT * FROM users
            WHERE username = %s
        '''
        g.cur.execute(sql, username)
        user = g.cur.fetchone()
        if user:
            flash("Account already exists")
            return redirect("/signup")

        sql = '''
            INSERT INTO users (username, password)
            VALUES (%s, %s)
        '''
        g.cur.execute(sql, (username.strip(), password))
        g.conn.commit()
        if g.cur.lastrowid:
            flash("Successful registration, please signin")
        else:
            flash("Unable to register at the moment, please try again later")

    return redirect("/signup")


@app.get("/login")
def render_login_page():
    if g.get("user"):
        return redirect("/")
    return render_template("login.html")


@app.post("/login")
def login():
    username = request.form.get("username")
    password = request.form.get("password")

    if not username or not password:
        flash("Username and password cannot be empty")
        return redirect("/login")
    else:
        sql = '''
            SELECT * FROM users
            WHERE username = %s
        '''
        g.cur.execute(sql, username)
        user = g.cur.fetchone()

        if not user or user["password"] != password:
            flash("Username or password are incorrect")
            return redirect("/login")
        else:
            response = make_response(redirect("/"))
            response.set_cookie("user_id", str(user["id"]))
            return response


@app.post("/logout")
def logout():
    response = make_response(redirect("/"))
    response.set_cookie("user_id", "")
    return response


@app.route("/", methods=["GET"])
def homepage():
    sql = '''
        SELECT * FROM posts
        JOIN users ON posts.author = users.id
        ORDER BY created DESC
    '''
    g.cur.execute(sql)

    posts = g.cur.fetchall()
    return render_template("index.html", posts=posts)


@app.route("/", methods=["POST"])
def new_post():
    if not g.user:
        flash("You must be signin")
        return redirect("/")

    title = request.form.get("title")
    content = request.form.get("content")

    if title and content:
        sql = '''
            INSERT INTO posts (title, content, author)
            VALUES (%s, %s, %s)
        '''
        g.cur.execute(sql, (title, content, g.user['id']))
        g.conn.commit()
        flash("Post created!")
    else:
        flash("Title and content cannot be empty")

    return redirect("/")


@app.route("/post/<post_id>", methods=["GET"])
def render_post(post_id):
    sql = '''
        SELECT * FROM posts
        WHERE id = %s
    '''
    g.cur.execute(sql, post_id)
    post = g.cur.fetchone()
    return render_template('post.html', post=post)


@app.route("/post/<post_id>", methods=["POST"])
def edit_post(post_id):
    title = request.form.get("title")
    content = request.form.get("content")

    if title and content:
        sql = '''
            SELECT * FROM posts
            WHERE id = %s
        '''
        g.cur.execute(sql, post_id)
        post = g.cur.fetchone()

        if not post:
            flash("Post does not exists")
        elif post['author'] != g.user['id']:
            flash("You cannot edit this post")
        else:
            sql = '''
                UPDATE posts
                SET title = %s,
                    content = %s
                WHERE
                    id = %s
            '''

            g.cur.execute(sql, (title, content, post_id))
            g.conn.commit()
            flash("Post updated!")

        return redirect(request.url)
    else:
        flash("Title and content cannot be empty")
        return redirect(request.url)


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/resignation-letter", methods=["GET"])
def render_form():
    return render_template("resignation-letter.html")


@app.route("/resignation-letter", methods=["POST"])
def create_file():
    fullname = request.form.get("fullname")
    reason = request.form.get("reason")

    if fullname and reason:
        doc = Document()
        doc.add_heading("Resignation Letter", level=0)
        doc.add_paragraph(fullname)
        doc.add_paragraph(reason)
        doc.save("./static/resignation-letter.docx")
        return render_template("resignation-letter.html", file_created=True)
    else:
        flash("Fullname and reason cannot be empty")
        return render_template("resignation-letter.html", file_created=False)


if __name__ == "__main__":
    app.run()
